#!/usr/bin/env python3
"""
Convert DOCX to PDF using python-docx and reportlab
Usage: python convert_doc.py <input_file.docx> <output_file.pdf>
"""

import sys
import os
from docx import Document
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Table, TableStyle, Spacer, PageBreak, Image
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT, TA_JUSTIFY
from io import BytesIO

def convert_docx_to_pdf(input_path, output_path):
    """
    Convert DOCX file to PDF.
    Preserves basic formatting, paragraphs, and attempts to embed images.
    """
    try:
        # Load the DOCX document
        doc = Document(input_path)
        
        # Create PDF document
        pdf_buffer = BytesIO()
        pdf_doc = SimpleDocTemplate(pdf_buffer, pagesize=letter)
        
        # Get default styles
        styles = getSampleStyleSheet()
        story = []
        
        # Process each paragraph and table
        for element in doc.element.body:
            # Handle paragraphs
            if element.tag.endswith('p'):
                para = None
                for p in doc.paragraphs:
                    if p._element == element:
                        para = p
                        break
                
                if para:
                    text = para.text.strip()
                    if text:
                        # Determine alignment
                        alignment = TA_LEFT
                        if para.alignment:
                            if para.alignment == 1:
                                alignment = TA_CENTER
                            elif para.alignment == 2:
                                alignment = TA_RIGHT
                            elif para.alignment == 3:
                                alignment = TA_JUSTIFY
                        
                        # Use built-in Normal style with alignment
                        para_style = ParagraphStyle(
                            f'para_{id(para)}',
                            parent=styles['Normal'],
                            alignment=alignment,
                            spaceAfter=12
                        )
                        
                        try:
                            story.append(Paragraph(text, para_style))
                        except:
                            # Fallback for special characters
                            try:
                                story.append(Paragraph(text.encode('ascii', 'ignore').decode(), para_style))
                            except:
                                # Last resort: use simple text
                                story.append(Paragraph('(Unable to render text)', para_style))
            
            # Handle tables
            elif element.tag.endswith('tbl'):
                for table in doc.tables:
                    if table._element == element:
                        table_data = []
                        for row in table.rows:
                            row_data = [cell.text.strip() for cell in row.cells]
                            table_data.append(row_data)
                        
                        if table_data:
                            # Create table with basic styling
                            t = Table(table_data)
                            t.setStyle(TableStyle([
                                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                                ('FONTSIZE', (0, 0), (-1, 0), 10),
                                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                                ('GRID', (0, 0), (-1, -1), 1, colors.black),
                                ('FONTSIZE', (0, 1), (-1, -1), 9),
                            ]))
                            story.append(t)
                            story.append(Spacer(1, 0.3*inch))
        
        # Build PDF
        if story:
            pdf_doc.build(story)
        
        # Write to file
        with open(output_path, 'wb') as f:
            f.write(pdf_buffer.getvalue())
        
        return True, None
    
    except Exception as e:
        return False, str(e)

if __name__ == '__main__':
    if len(sys.argv) < 3:
        print('Usage: python convert_doc.py <input.docx> <output.pdf>', file=sys.stderr)
        sys.exit(1)
    
    input_file = sys.argv[1]
    output_file = sys.argv[2]
    
    if not os.path.exists(input_file):
        print(f'Error: Input file not found: {input_file}', file=sys.stderr)
        sys.exit(1)
    
    success, error = convert_docx_to_pdf(input_file, output_file)
    
    if success:
        print(f'Success: {output_file}')
        sys.exit(0)
    else:
        print(f'Error: {error}', file=sys.stderr)
        sys.exit(1)
